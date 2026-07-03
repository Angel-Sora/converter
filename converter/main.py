# Universal File Converter
# Copyright (c) 2026 Angel-Sora
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
import os
import sys
import json
import csv
import pandas as pd
from pathlib import Path
from datetime import datetime
import logging
import shutil
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import docx
from docx import Document
import openpyxl
import xml.etree.ElementTree as ET
import yaml
import chardet
from concurrent.futures import ThreadPoolExecutor, as_completed
import argparse
from tqdm import tqdm
import zipfile
import hashlib

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('file_converter.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class FileConverter:
    """
    Универсальный конвертер файлов
    Поддерживает: изображения, документы, таблицы, данные, архивы
    """
    
    SUPPORTED_FORMATS = {
        # Изображения
        'image': {
            'input': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.gif'],
            'output': ['.jpg', '.png', '.bmp', '.webp', '.pdf']
        },
        # Документы
        'document': {
            'input': ['.txt', '.docx', '.pdf', '.rtf', '.odt'],
            'output': ['.txt', '.docx', '.pdf', '.html', '.md']
        },
        # Таблицы
        'spreadsheet': {
            'input': ['.csv', '.xlsx', '.xls', '.ods', '.json', '.xml'],
            'output': ['.csv', '.xlsx', '.json', '.html', '.sqlite']
        },
        # Данные
        'data': {
            'input': ['.json', '.xml', '.yaml', '.yml', '.toml'],
            'output': ['.json', '.xml', '.yaml', '.csv', '.pickle']
        },
        # Архивы
        'archive': {
            'input': ['.zip', '.tar', '.gz', '.rar'],
            'output': ['.zip', '.tar', '.gz']
        }
    }
    
    def __init__(self):
        self.conversion_stats = {
            'success': 0,
            'failed': 0,
            'skipped': 0,
            'total_size_before': 0,
            'total_size_after': 0
        }
        self.conversion_history = []
        self.output_dir = None
    
    # ============ ОСНОВНАЯ ЛОГИКА ============
    
    def convert_file(self, input_path: str, output_path: str = None, 
                    output_format: str = None, **kwargs) -> bool:
        """
        Конвертация одного файла
        """
        input_path = Path(input_path)
        
        if not input_path.exists():
            logging.error(f"❌ Файл не найден: {input_path}")
            return False
        
        # Определяем формат вывода
        if output_format:
            output_format = output_format.lower()
            if not output_format.startswith('.'):
                output_format = '.' + output_format
        elif output_path:
            output_path = Path(output_path)
            output_format = output_path.suffix.lower()
        else:
            output_format = self._suggest_format(input_path.suffix.lower())
            if not output_format:
                logging.error(f"❌ Не удалось определить формат вывода")
                return False
        
        # Если output_path не указан, создаем автоматически
        if not output_path:
            base_name = input_path.stem
            if self.output_dir:
                output_dir = Path(self.output_dir)
            else:
                output_dir = input_path.parent / 'converted'
            
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"{base_name}{output_format}"
        else:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Запоминаем размер
        size_before = input_path.stat().st_size
        
        try:
            # Определяем тип файла и вызываем соответствующий метод
            input_ext = input_path.suffix.lower()
            
            # Определяем категорию файла
            category = self._get_category(input_ext)
            
            if category == 'image':
                success = self._convert_image(input_path, output_path, **kwargs)
            elif category == 'document':
                success = self._convert_document(input_path, output_path, **kwargs)
            elif category == 'spreadsheet':
                success = self._convert_spreadsheet(input_path, output_path, **kwargs)
            elif category == 'data':
                success = self._convert_data(input_path, output_path, **kwargs)
            elif category == 'archive':
                success = self._convert_archive(input_path, output_path, **kwargs)
            else:
                logging.error(f"❌ Неподдерживаемый формат: {input_ext}")
                return False
            
            if success:
                # Обновляем статистику
                size_after = output_path.stat().st_size
                self.conversion_stats['success'] += 1
                self.conversion_stats['total_size_before'] += size_before
                self.conversion_stats['total_size_after'] += size_after
                
                self.conversion_history.append({
                    'input': str(input_path),
                    'output': str(output_path),
                    'status': 'success',
                    'time': datetime.now().isoformat()
                })
                
                compression = (1 - size_after / size_before) * 100 if size_before > 0 else 0
                logging.info(f"✅ Конвертировано: {input_path.name} → {output_path.name} "
                           f"({size_before:,} → {size_after:,} байт, {compression:.1f}%)")
                return True
            else:
                self.conversion_stats['failed'] += 1
                logging.error(f"❌ Ошибка конвертации: {input_path}")
                return False
                
        except Exception as e:
            self.conversion_stats['failed'] += 1
            logging.error(f"❌ Ошибка при конвертации {input_path}: {e}")
            return False
    
    def convert_batch(self, input_dir: str, output_dir: str = None, 
                     output_format: str = None, patterns: list = None,
                     recursive: bool = False, max_workers: int = 4) -> dict:
        """
        Пакетная конвертация файлов
        """
        input_dir = Path(input_dir)
        
        if not input_dir.exists():
            logging.error(f"❌ Папка не найдена: {input_dir}")
            return self.conversion_stats
        
        # Устанавливаем выходную папку
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            self.output_dir = input_dir / 'converted'
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Собираем файлы для конвертации
        files_to_convert = []
        
        if recursive:
            all_files = list(input_dir.rglob('*'))
        else:
            all_files = list(input_dir.glob('*'))
        
        for file_path in all_files:
            if not file_path.is_file():
                continue
            
            # Проверяем паттерны
            if patterns:
                if not any(file_path.name.endswith(p) for p in patterns):
                    continue
            
            files_to_convert.append(file_path)
        
        if not files_to_convert:
            logging.warning("⚠️ Нет файлов для конвертации")
            return self.conversion_stats
        
        logging.info(f"📋 Найдено файлов для конвертации: {len(files_to_convert)}")
        
        # Конвертируем с прогресс-баром
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = []
            for file_path in files_to_convert:
                future = executor.submit(
                    self.convert_file, 
                    str(file_path), 
                    None, 
                    output_format
                )
                futures.append(future)
            
            # Прогресс-бар
            for future in tqdm(as_completed(futures), total=len(futures), desc="Конвертация"):
                pass
        
        # Генерируем отчет
        self._generate_report()
        
        return self.conversion_stats
    
    # ============ КОНВЕРТЕРЫ ПО КАТЕГОРИЯМ ============
    
    def _convert_image(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Конвертация изображений"""
        try:
            img = Image.open(input_path)
            
            # Опции конвертации
            quality = kwargs.get('quality', 95)
            resize = kwargs.get('resize')
            
            if resize:
                img = img.resize(resize, Image.Resampling.LANCZOS)
            
            # Сохраняем в нужном формате
            output_ext = output_path.suffix.lower()
            
            if output_ext in ['.jpg', '.jpeg']:
                img.save(output_path, 'JPEG', quality=quality)
            elif output_ext == '.png':
                img.save(output_path, 'PNG')
            elif output_ext == '.bmp':
                img.save(output_path, 'BMP')
            elif output_ext == '.webp':
                img.save(output_path, 'WEBP', quality=quality)
            elif output_ext == '.pdf':
                # Конвертация в PDF (для одного изображения)
                img.save(output_path, 'PDF', resolution=100.0)
            else:
                img.save(output_path)
            
            return True
        except Exception as e:
            logging.error(f"Ошибка конвертации изображения: {e}")
            return False
    
    def _convert_document(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Конвертация документов"""
        input_ext = input_path.suffix.lower()
        output_ext = output_path.suffix.lower()
        
        try:
            if input_ext == '.txt':
                # Читаем текстовый файл с определением кодировки
                with open(input_path, 'rb') as f:
                    raw_data = f.read()
                    encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                
                content = raw_data.decode(encoding)
                
                if output_ext == '.docx':
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(output_path)
                    return True
                elif output_ext == '.html':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(f"<html><body><pre>{content}</pre></body></html>")
                    return True
                elif output_ext == '.md':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    return True
            
            elif input_ext == '.docx':
                doc = Document(input_path)
                content = '\n'.join([p.text for p in doc.paragraphs])
                
                if output_ext == '.txt':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    return True
                elif output_ext == '.html':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        html_content = ''.join([f"<p>{p.text}</p>" for p in doc.paragraphs])
                        f.write(f"<html><body>{html_content}</body></html>")
                    return True
            
            elif input_ext == '.pdf' and output_ext == '.txt':
                # OCR для PDF
                images = convert_from_path(input_path)
                text = ''
                for image in images:
                    text += pytesseract.image_to_string(image, lang='rus+eng')
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return True
            
            else:
                logging.warning(f"⚠️ Конвертация {input_ext} → {output_ext} не поддерживается")
                return False
                
        except Exception as e:
            logging.error(f"Ошибка конвертации документа: {e}")
            return False
    
    def _convert_spreadsheet(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Конвертация таблиц"""
        input_ext = input_path.suffix.lower()
        output_ext = output_path.suffix.lower()
        
        try:
            # Читаем данные в DataFrame
            if input_ext == '.csv':
                # Определяем кодировку
                with open(input_path, 'rb') as f:
                    raw = f.read()
                    encoding = chardet.detect(raw)['encoding'] or 'utf-8'
                
                df = pd.read_csv(input_path, encoding=encoding)
            elif input_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(input_path)
            elif input_ext == '.json':
                df = pd.read_json(input_path)
            elif input_ext == '.xml':
                # Простой парсинг XML
                tree = ET.parse(input_path)
                root = tree.getroot()
                data = []
                for child in root:
                    row = {}
                    for subchild in child:
                        row[subchild.tag] = subchild.text
                    data.append(row)
                df = pd.DataFrame(data)
            else:
                logging.warning(f"⚠️ Неподдерживаемый формат таблицы: {input_ext}")
                return False
            
            # Сохраняем в нужном формате
            if output_ext == '.csv':
                df.to_csv(output_path, index=False, encoding='utf-8-sig')
            elif output_ext in ['.xlsx', '.xls']:
                df.to_excel(output_path, index=False)
            elif output_ext == '.json':
                df.to_json(output_path, orient='records', force_ascii=False)
            elif output_ext == '.html':
                df.to_html(output_path, index=False)
            elif output_ext == '.sqlite':
                import sqlite3
                conn = sqlite3.connect(output_path)
                df.to_sql('data', conn, if_exists='replace', index=False)
                conn.close()
            else:
                logging.warning(f"⚠️ Конвертация {input_ext} → {output_ext} не поддерживается")
                return False
            
            return True
            
        except Exception as e:
            logging.error(f"Ошибка конвертации таблицы: {e}")
            return False
    
    def _convert_data(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Конвертация данных (JSON, YAML, XML и т.д.)"""
        input_ext = input_path.suffix.lower()
        output_ext = output_path.suffix.lower()
        
        try:
            # Читаем данные
            if input_ext == '.json':
                with open(input_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            elif input_ext in ['.yaml', '.yml']:
                with open(input_path, 'r', encoding='utf-8') as f:
                    data = yaml.safe_load(f)
            elif input_ext == '.xml':
                tree = ET.parse(input_path)
                root = tree.getroot()
                data = self._xml_to_dict(root)
            else:
                logging.warning(f"⚠️ Неподдерживаемый формат данных: {input_ext}")
                return False
            
            # Сохраняем в нужном формате
            if output_ext == '.json':
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            elif output_ext in ['.yaml', '.yml']:
                with open(output_path, 'w', encoding='utf-8') as f:
                    yaml.dump(data, f, allow_unicode=True, default_flow_style=False)
            elif output_ext == '.xml':
                root = self._dict_to_xml('root', data)
                tree = ET.ElementTree(root)
                tree.write(output_path, encoding='utf-8', xml_declaration=True)
            elif output_ext == '.csv' and isinstance(data, list):
                df = pd.DataFrame(data)
                df.to_csv(output_path, index=False, encoding='utf-8-sig')
            else:
                logging.warning(f"⚠️ Конвертация {input_ext} → {output_ext} не поддерживается")
                return False
            
            return True
            
        except Exception as e:
            logging.error(f"Ошибка конвертации данных: {e}")
            return False
    
    def _convert_archive(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Конвертация архивов (распаковка/упаковка)"""
        input_ext = input_path.suffix.lower()
        output_ext = output_path.suffix.lower()
        
        try:
            if input_ext == '.zip':
                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                    if output_ext == '.zip':
                        # Просто копируем
                        shutil.copy2(input_path, output_path)
                    else:
                        # Распаковываем в папку
                        extract_dir = output_path.with_suffix('')
                        zip_ref.extractall(extract_dir)
                        logging.info(f"📦 Распакован архив в: {extract_dir}")
                    return True
            else:
                logging.warning(f"⚠️ Работа с архивами {input_ext} → {output_ext} ограничена")
                return False
                
        except Exception as e:
            logging.error(f"Ошибка работы с архивом: {e}")
            return False
    
    # ============ ВСПОМОГАТЕЛЬНЫЕ МЕТОДЫ ============
    
    def _get_category(self, ext: str) -> str:
        """Определение категории файла по расширению"""
        for category, formats in self.SUPPORTED_FORMATS.items():
            if ext in formats['input']:
                return category
        return 'unknown'
    
    def _suggest_format(self, ext: str) -> str:
        """Предложение формата вывода на основе входного"""
        suggestions = {
            '.txt': '.docx',
            '.docx': '.txt',
            '.pdf': '.txt',
            '.jpg': '.png',
            '.png': '.jpg',
            '.csv': '.xlsx',
            '.xlsx': '.csv',
            '.json': '.csv',
            '.xml': '.json'
        }
        return suggestions.get(ext, None)
    
    def _xml_to_dict(self, element) -> dict:
        """Преобразование XML элемента в словарь"""
        result = {}
        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text
            else:
                result[child.tag] = self._xml_to_dict(child)
        return result
    
    def _dict_to_xml(self, tag: str, data) -> ET.Element:
        """Преобразование словаря в XML элемент"""
        element = ET.Element(tag)
        if isinstance(data, dict):
            for key, value in data.items():
                child = self._dict_to_xml(key, value)
                element.append(child)
        elif isinstance(data, list):
            for item in data:
                child = self._dict_to_xml('item', item)
                element.append(child)
        else:
            element.text = str(data)
        return element
    
    def _generate_report(self) -> None:
        """Генерация отчета о конвертации"""
        report_path = self.output_dir / 'conversion_report.txt'
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("📊 ОТЧЕТ О КОНВЕРТАЦИИ ФАЙЛОВ\n")
            f.write("="*80 + "\n")
            f.write(f"📅 Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"✅ Успешно: {self.conversion_stats['success']}\n")
            f.write(f"❌ Ошибок: {self.conversion_stats['failed']}\n")
            f.write(f"⏭️ Пропущено: {self.conversion_stats['skipped']}\n")
            f.write(f"\n📊 Размеры:\n")
            f.write(f"   До: {self.conversion_stats['total_size_before']:,} байт\n")
            f.write(f"   После: {self.conversion_stats['total_size_after']:,} байт\n")
            if self.conversion_stats['total_size_before'] > 0:
                compression = (1 - self.conversion_stats['total_size_after'] / 
                             self.conversion_stats['total_size_before']) * 100
                f.write(f"   Сжатие: {compression:.1f}%\n")
            f.write("\n📋 История конвертаций:\n")
            for item in self.conversion_history:
                f.write(f"   {item['time']} - {item['input']} → {item['output']}\n")
            f.write("="*80 + "\n")
        
        logging.info(f"📄 Отчет сохранен: {report_path}")


# ============ GUI ИНТЕРФЕЙС ============

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

class ConverterGUI:
    """Графический интерфейс для конвертера"""
    
    def __init__(self):
        self.converter = FileConverter()
        self.root = tk.Tk()
        self.root.title("🔄 Универсальный конвертер файлов")
        self.root.geometry("800x600")
        
        self.setup_ui()
        
    def setup_ui(self):
        """Настройка интерфейса"""
        # Главный фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Заголовок
        title = ttk.Label(main_frame, text="🔄 Универсальный конвертер файлов", 
                         font=('Arial', 16, 'bold'))
        title.grid(row=0, column=0, columnspan=3, pady=10)
        
        # Выбор файлов
        ttk.Label(main_frame, text="📁 Файлы для конвертации:").grid(row=1, column=0, sticky=tk.W)
        
        self.file_listbox = tk.Listbox(main_frame, height=8, selectmode=tk.EXTENDED)
        self.file_listbox.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Кнопки для работы с файлами
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=3, column=0, columnspan=3, pady=5)
        
        ttk.Button(btn_frame, text="➕ Добавить файлы", 
                  command=self.add_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="📂 Добавить папку", 
                  command=self.add_folder).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🗑️ Удалить выбранные", 
                  command=self.remove_selected).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🧹 Очистить всё", 
                  command=self.clear_all).pack(side=tk.LEFT, padx=5)
        
        # Настройки конвертации
        settings_frame = ttk.LabelFrame(main_frame, text="⚙️ Настройки", padding="10")
        settings_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        # Формат вывода
        ttk.Label(settings_frame, text="Формат вывода:").grid(row=0, column=0, sticky=tk.W)
        self.output_format_var = tk.StringVar()
        self.output_format_combo = ttk.Combobox(settings_frame, textvariable=self.output_format_var,
                                               values=['.jpg', '.png', '.pdf', '.txt', '.docx', 
                                                      '.csv', '.xlsx', '.json', '.html'])
        self.output_format_combo.grid(row=0, column=1, padx=10, sticky=tk.W)
        self.output_format_combo.set('.jpg')
        
        # Выходная папка
        ttk.Label(settings_frame, text="Выходная папка:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.output_dir_var = tk.StringVar(value="./converted")
        ttk.Entry(settings_frame, textvariable=self.output_dir_var, width=40).grid(row=1, column=1, padx=10)
        ttk.Button(settings_frame, text="📁 Выбрать", 
                  command=self.select_output_dir).grid(row=1, column=2)
        
        # Дополнительные опции
        self.resize_var = tk.BooleanVar()
        ttk.Checkbutton(settings_frame, text="Изменить размер изображений", 
                       variable=self.resize_var).grid(row=2, column=0, columnspan=2, sticky=tk.W)
        
        self.recursive_var = tk.BooleanVar()
        ttk.Checkbutton(settings_frame, text="Рекурсивный обход папок", 
                       variable=self.recursive_var).grid(row=3, column=0, columnspan=2, sticky=tk.W)
        
        # Кнопка запуска
        ttk.Button(main_frame, text="🚀 НАЧАТЬ КОНВЕРТАЦИЮ", 
                  command=self.start_conversion,
                  style='Accent.TButton').grid(row=5, column=0, columnspan=3, pady=15)
        
        # Лог
        log_frame = ttk.LabelFrame(main_frame, text="📝 Лог", padding="5")
        log_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=8, width=70)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # Настройка веса
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=1)
        
        # Перенаправление логов в GUI
        class TextHandler(logging.Handler):
            def __init__(self, text_widget):
                logging.Handler.__init__(self)
                self.text_widget = text_widget
            
            def emit(self, record):
                msg = self.format(record)
                self.text_widget.insert(tk.END, msg + '\n')
                self.text_widget.see(tk.END)
        
        text_handler = TextHandler(self.log_text)
        text_handler.setLevel(logging.INFO)
        text_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        logging.getLogger().addHandler(text_handler)
    
    def add_files(self):
        """Добавление файлов"""
        files = filedialog.askopenfilenames()
        for f in files:
            self.file_listbox.insert(tk.END, f)
    
    def add_folder(self):
        """Добавление папки"""
        folder = filedialog.askdirectory()
        if folder:
            for file in Path(folder).rglob('*'):
                if file.is_file():
                    self.file_listbox.insert(tk.END, str(file))
    
    def remove_selected(self):
        """Удаление выбранных файлов"""
        selected = self.file_listbox.curselection()
        for i in reversed(selected):
            self.file_listbox.delete(i)
    
    def clear_all(self):
        """Очистка списка"""
        self.file_listbox.delete(0, tk.END)
    
    def select_output_dir(self):
        """Выбор выходной папки"""
        dir_path = filedialog.askdirectory()
        if dir_path:
            self.output_dir_var.set(dir_path)
    
    def start_conversion(self):
        """Запуск конвертации"""
        files = list(self.file_listbox.get(0, tk.END))
        
        if not files:
            messagebox.showwarning("Предупреждение", "Добавьте файлы для конвертации!")
            return
        
        output_format = self.output_format_var.get()
        output_dir = self.output_dir_var.get()
        
        # Устанавливаем выходную папку
        self.converter.output_dir = Path(output_dir)
        
        # Запускаем конвертацию
        success_count = 0
        for file_path in files:
            if self.converter.convert_file(file_path, output_format=output_format):
                success_count += 1
        
        messagebox.showinfo("Готово", f"✅ Конвертация завершена!\nУспешно: {success_count} из {len(files)}")
    
    def run(self):
        """Запуск GUI"""
        self.root.mainloop()


# ============ КОМАНДНАЯ СТРОКА ============

def main_cli():
    """Интерфейс командной строки"""
    parser = argparse.ArgumentParser(description='🔄 Универсальный конвертер файлов')
    parser.add_argument('input', help='Входной файл или папка')
    parser.add_argument('-o', '--output', help='Выходной файл или папка')
    parser.add_argument('-f', '--format', help='Формат вывода (например: .jpg)')
    parser.add_argument('-r', '--recursive', action='store_true', help='Рекурсивный обход папок')
    parser.add_argument('-w', '--workers', type=int, default=4, help='Количество потоков')
    parser.add_argument('--gui', action='store_true', help='Запустить GUI')
    
    args = parser.parse_args()
    
    if args.gui:
        app = ConverterGUI()
        app.run()
        return
    
    converter = FileConverter()
    
    input_path = Path(args.input)
    
    if input_path.is_file():
        # Конвертация одного файла
        output_path = Path(args.output) if args.output else None
        converter.convert_file(str(input_path), output_path, args.format)
    elif input_path.is_dir():
        # Пакетная конвертация
        converter.convert_batch(
            str(input_path), 
            args.output, 
            args.format,
            recursive=args.recursive,
            max_workers=args.workers
        )
    else:
        print(f"❌ Не найден: {input_path}")

# ============ ТОЧКА ВХОДА ============

if __name__ == "__main__":
    if len(sys.argv) > 1:
        main_cli()
    else:
        # Если нет аргументов - запускаем GUI
        app = ConverterGUI()
        app.run()
